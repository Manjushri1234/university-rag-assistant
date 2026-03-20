import os
import pandas as pd
from langchain_community.document_loaders import PyPDFLoader
from docx import Document
from langchain_core.documents import Document as LangDocument


# -------------------------------
# Load Word (.docx)
# -------------------------------
def load_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
    
    return [LangDocument(page_content=text, metadata={"source": file_path})]


# -------------------------------
# Load PDF (.pdf)
# -------------------------------
def load_pdf(file_path):
    loader = PyPDFLoader(file_path)
    return loader.load()


# -------------------------------
# Load Excel (.xlsx)
# -------------------------------
def load_excel(file_path):
    df = pd.read_excel(file_path)

    documents = []
    for _, row in df.iterrows():
        text = ", ".join([f"{col}: {row[col]}" for col in df.columns])
        documents.append(
            LangDocument(page_content=text, metadata={"source": file_path})
        )

    return documents


# -------------------------------
# Load ALL files
# -------------------------------
def load_all_documents(data_path="data"):
    documents = []

    for file in os.listdir(data_path):
        file_path = os.path.join(data_path, file)

        if file.endswith(".pdf"):
            documents.extend(load_pdf(file_path))

        elif file.endswith(".docx"):
            documents.extend(load_docx(file_path))

        elif file.endswith(".xlsx"):
            documents.extend(load_excel(file_path))

    return documents
